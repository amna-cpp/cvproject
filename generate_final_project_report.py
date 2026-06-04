from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


PROJECT_ROOT = Path(__file__).resolve().parent
IMPLEMENTATION_DIR = PROJECT_ROOT / "implementation"
RESULTS_DIR = IMPLEMENTATION_DIR / "results"
METRICS_DIR = RESULTS_DIR / "metrics"
PLOTS_DIR = RESULTS_DIR / "plots"
OUTPUT_DIR = PROJECT_ROOT / "final_report"
OUTPUT_PDF = OUTPUT_DIR / "Final_Project_Report.pdf"
OUTPUT_DOCX = OUTPUT_DIR / "Final_Project_Report.docx"

FILTER_ORDER = ["original", "blur", "brightness", "skin_smooth", "eye_enlarge", "face_slim", "color_tone"]
FILTERS = [f for f in FILTER_ORDER if f != "original"]

CONFIG = {
    "project_title": "Facial Filter Impact on Face Recognition",
    "course": "Computer Vision, Spring 2026",
    "institution": "National University of Computer and Emerging Sciences, FAST School of Computing",
    "paper": "A Comprehensive Evaluation Framework for the Study of the Effects of Facial Filters on Face Recognition Accuracy",
    "paper_url": "https://arxiv.org/abs/2507.17729",
    "dataset": "Labeled Faces in the Wild (LFW)",
    "subset_pairs": 600,
    "genuine_pairs": 300,
    "impostor_pairs": 300,
    "unique_pair_images": 807,
    "original_images": 13233,
    "identities": 5749,
    "filtered_images": 4842,
}


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def safe_float(value: object, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return default


def load_results() -> dict:
    arcface = read_json(METRICS_DIR / "recognition_metrics.json")
    improved = read_json(METRICS_DIR / "recognition_metrics_adaface.json")
    mitigation = read_json(RESULTS_DIR / "mitigation_linear" / "linear_mitigation_metrics.json")
    summary = pd.read_csv(PLOTS_DIR / "comparison" / "summary_table.csv")
    return {
        "arcface": arcface,
        "improved": improved,
        "mitigation": mitigation,
        "summary": summary,
    }


def avg(values: list[float]) -> float:
    return sum(values) / len(values) if values else 0.0


def metrics_summary(results: dict) -> dict:
    arc = results["arcface"]
    imp = results["improved"]
    table = results["summary"]
    arc_filter_acc = avg([arc[f]["accuracy"] for f in FILTERS])
    arc_filter_fnmr = avg([arc[f]["fnmr_fmr01"] for f in FILTERS])
    imp_filter_acc = avg([imp[f]["accuracy"] for f in FILTERS])
    imp_filter_fnmr = avg([imp[f]["fnmr_fmr01"] for f in FILTERS])
    paper_fnmr = safe_float(table["Paper FNMR"].mean())
    improved_fnmr = safe_float(table["AdaFace"].mean())
    linear_fnmr = safe_float(table["+LinearCorr"].mean())
    meanshift_fnmr = safe_float(table["+MeanShift"].mean())
    return {
        "arc_original_acc": arc["original"]["accuracy"],
        "imp_original_acc": imp["original"]["accuracy"],
        "arc_filter_acc": arc_filter_acc,
        "arc_filter_fnmr": arc_filter_fnmr,
        "imp_filter_acc": imp_filter_acc,
        "imp_filter_fnmr": imp_filter_fnmr,
        "meanshift_fnmr": meanshift_fnmr,
        "linear_fnmr": linear_fnmr,
        "improved_reduction": ((paper_fnmr - improved_fnmr) / paper_fnmr * 100) if paper_fnmr else 0,
        "combined_reduction": ((paper_fnmr - linear_fnmr) / paper_fnmr * 100) if paper_fnmr else 0,
    }


def recognition_rows(metrics: dict) -> list[list[str]]:
    rows = [["Filter", "Accuracy", "FNMR@FMR=1%", "FNMR@FMR=0.1%"]]
    for f in FILTER_ORDER:
        m = metrics[f]
        rows.append([f, pct(m["accuracy"]), pct(m["fnmr_fmr01"]), pct(m["fnmr_fmr001"])])
    return rows


def comparison_rows(summary: pd.DataFrame) -> list[list[str]]:
    rows = [["Filter", "ArcFace FNMR", "Mean-shift", "Improved Backbone", "Linear Corrector"]]
    for _, row in summary.iterrows():
        rows.append([
            str(row["Filter"]),
            pct(safe_float(row["Paper FNMR"])),
            pct(safe_float(row["+MeanShift"])),
            pct(safe_float(row["AdaFace"])),
            pct(safe_float(row["+LinearCorr"])),
        ])
    return rows


def build_pdf(results: dict) -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    summary = metrics_summary(results)
    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("TitleCenter", parent=styles["Title"], alignment=TA_CENTER, fontSize=20, leading=24, textColor=colors.HexColor("#0B2545")))
    styles.add(ParagraphStyle("SubtitleCenter", parent=styles["Normal"], alignment=TA_CENTER, fontSize=11, leading=15, textColor=colors.HexColor("#1F4D78")))
    styles.add(ParagraphStyle("H1Custom", parent=styles["Heading1"], fontSize=15, leading=19, textColor=colors.HexColor("#185FA5"), spaceBefore=14, spaceAfter=8))
    styles.add(ParagraphStyle("H2Custom", parent=styles["Heading2"], fontSize=12, leading=15, textColor=colors.HexColor("#0C447C"), spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle("BodyJustify", parent=styles["BodyText"], alignment=TA_JUSTIFY, fontSize=9.6, leading=13.5, spaceAfter=6))
    styles.add(ParagraphStyle("BodySmall", parent=styles["BodyText"], alignment=TA_LEFT, fontSize=8.8, leading=11.5, spaceAfter=4))
    styles.add(ParagraphStyle("Caption", parent=styles["BodyText"], alignment=TA_CENTER, fontSize=8.5, leading=10.5, textColor=colors.HexColor("#555555"), spaceBefore=2, spaceAfter=8))

    story: list = []

    def h1(text: str) -> None:
        story.append(Paragraph(text, styles["H1Custom"]))

    def h2(text: str) -> None:
        story.append(Paragraph(text, styles["H2Custom"]))

    def p(text: str) -> None:
        story.append(Paragraph(text, styles["BodyJustify"]))

    def bullet(text: str) -> None:
        story.append(Paragraph(f"- {text}", styles["BodyJustify"]))

    def table(data: list[list[str]], widths: list[float] | None = None) -> None:
        if widths is None:
            widths = [doc.width / len(data[0])] * len(data[0])
        t = Table(data, colWidths=widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6F1FB")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.0),
            ("LEADING", (0, 0), (-1, -1), 9.5),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B5D4F4")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F8F8")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.22 * cm))

    def plot(rel: str, caption: str) -> None:
        path = PLOTS_DIR / rel
        if path.exists():
            img = Image(str(path))
            max_w = doc.width
            max_h = 9.8 * cm
            scale = min(max_w / img.drawWidth, max_h / img.drawHeight)
            img.drawWidth *= scale
            img.drawHeight *= scale
            story.append(img)
            story.append(Paragraph(caption, styles["Caption"]))
        else:
            story.append(Paragraph(f"Plot missing: {rel}", styles["Caption"]))

    story.append(Spacer(1, 2.2 * cm))
    story.append(Paragraph(CONFIG["project_title"], styles["TitleCenter"]))
    story.append(Paragraph("Final Project Report", styles["SubtitleCenter"]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(CONFIG["institution"], styles["SubtitleCenter"]))
    story.append(Paragraph(CONFIG["course"], styles["SubtitleCenter"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(f"Based on: {CONFIG['paper']}", styles["SubtitleCenter"]))
    story.append(Paragraph("Ozturk, Conwill, Gutierrez, Bowyer, and Scheirer", styles["SubtitleCenter"]))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", styles["SubtitleCenter"]))
    story.append(PageBreak())

    h1("Abstract")
    p(
        "This project implements an end-to-end computer vision system for evaluating how facial filters affect face verification accuracy. "
        "The work reimplements the core methodology of Ozturk et al. by preparing verification pairs, applying facial filter transformations, extracting face embeddings, computing biometric recognition metrics, and demonstrating the system through a live Gradio application. "
        "The prototype uses the LFW dataset and six programmatic filter categories that approximate softening, color grading, skin smoothing, eye enlargement, face slimming, and warm-tone transformations. "
        f"On the submitted 600-pair subset, the ArcFace baseline achieved an average filtered accuracy of {pct(summary['arc_filter_acc'])}, while the improved backbone path achieved {pct(summary['imp_filter_acc'])}. "
        f"The improved backbone reduced average FNMR@FMR=1% by {summary['improved_reduction']:.1f}% relative to the ArcFace baseline. "
        "The project also evaluates embedding-space mitigation using mean-shift and a Ridge-regression linear corrector, showing that learned correction can help on average but remains sensitive to filter type and calibration size."
    )

    h1("1. Introduction and Problem Definition")
    p(
        "Face recognition systems are increasingly exposed to images modified by social media and camera-app filters. "
        "These filters can change texture, color, illumination, and even local facial geometry. While many filters are visually subtle to humans, they may shift the embedding produced by a face recognition model and reduce verification accuracy. "
        "The problem addressed in this project is to measure, visualize, and mitigate the recognition degradation caused by such filters in a reproducible computer vision pipeline."
    )
    p(
        "The real-world challenge is important because face verification is used in authentication, account recovery, identity checks, and digital media workflows. "
        "If filtered photos are accepted as ordinary face images without evaluation, the system may reject genuine users or behave inconsistently across filter types. "
        "The goal of this project is not to propose a completely new recognition algorithm; it is to reproduce the evaluation framework from the selected paper and implement a working prototype that can demonstrate the full pipeline from data preparation to live demo."
    )

    h1("2. Literature Review")
    h2("2.1 Facial Filter Evaluation Framework")
    p(
        "Ozturk et al. introduced a structured framework for studying the effects of social media AR facial filters on face recognition. "
        "Their framework combines a controlled dataset, a principled filter selection process, and biometric evaluation metrics such as FNMR at fixed FMR thresholds and DET curves. "
        "The original study evaluated 125 filters from Instagram, Snapchat, Meitu, and Pitu, showing that structural filters such as eye enlargement and face slimming are more disruptive than simple color or brightness filters."
    )
    h2("2.2 Recognition Backbones")
    p(
        "ArcFace is a widely used face recognition method that learns discriminative angular embeddings through an additive angular margin loss. "
        "It provides strong baseline performance, but it does not explicitly adapt the margin to image quality. AdaFace later introduced a quality-adaptive margin that uses feature norm as a proxy for image quality, making it relevant for filter-degraded or low-quality face images. "
        "In this implementation, the code attempts to use AdaFace; because AdaFace was unavailable in the submitted Python 3.13/macOS environment, the improved-backbone path uses DeepFace GhostFaceNet as a 512-dimensional fallback."
    )
    h2("2.3 Dataset and Evaluation Protocols")
    p(
        "The Labeled Faces in the Wild dataset is a public benchmark for unconstrained face recognition. "
        "It provides identity folders and a standard pair protocol for genuine and impostor verification experiments. "
        "This project uses LFW instead of FRGC because LFW is publicly accessible and suitable for a course project demonstration. "
        "The evaluation still follows the selected paper's biometric protocol: cosine similarity is computed between embeddings, thresholds are swept, and accuracy, FNMR, FMR, and DET curves are reported."
    )

    h1("3. Methodology and System Architecture")
    p(
        "The system follows a modular pipeline: dataset setup, filter application, filter impact scoring, embedding extraction, recognition evaluation, mitigation, comparison visualization, and deployment through a web demo. "
        "Every stage is implemented as a runnable Python script and can be executed independently or through the master pipeline."
    )
    table([
        ["Stage", "Script / Module", "Purpose"],
        ["Dataset setup", "src/01_setup_dataset.py", "Extract LFW, parse standard pairs, save genuine and impostor CSVs."],
        ["Filter simulation", "src/02_apply_filters.py", "Generate six filter categories and preserve identity folder structure."],
        ["Filter ranking", "src/03_filter_selection.py", "Measure embedding displacement between original and filtered images."],
        ["Embedding extraction", "src/04_extract_embeddings.py, src/04b_extract_embeddings_adaface.py", "Extract and cache 512-D embeddings for original and filtered images."],
        ["Evaluation", "src/05_evaluate_recognition.py", "Compute cosine similarity, accuracy, FNMR, FMR, and DET curves."],
        ["Mitigation", "src/06_mitigation.py, src/06b_mitigation_linear.py", "Train filter detection, mean-shift correction, and Ridge linear correctors."],
        ["Deployment", "app/app.py", "Provide a live Gradio dashboard and verification demo."],
    ], [2.5 * cm, 4.5 * cm, 9.0 * cm])
    h2("Evaluation Metrics")
    p(
        "For each pair of embeddings e1 and e2, cosine similarity is computed as the dot product divided by the product of vector norms. "
        "For a threshold t, an impostor pair is falsely accepted when its similarity is at least t, producing FMR. "
        "A genuine pair is falsely rejected when its similarity is below t, producing FNMR. "
        "The report emphasizes accuracy at the best threshold and FNMR at FMR=1%, matching common biometric evaluation practice."
    )

    h1("4. Dataset Description")
    p(
        f"The dataset is based on {CONFIG['dataset']}. The prepared original data contains {CONFIG['original_images']:,} images across {CONFIG['identities']:,} identities. "
        f"For the submitted experiment, {CONFIG['genuine_pairs']} genuine and {CONFIG['impostor_pairs']} impostor pairs are used, producing {CONFIG['subset_pairs']} evaluation pairs. "
        f"The pair subset references {CONFIG['unique_pair_images']} unique images. Each of those images is transformed under six filter categories, producing {CONFIG['filtered_images']:,} generated filtered images."
    )
    table([
        ["Dataset Item", "Value"],
        ["Original LFW identities", f"{CONFIG['identities']:,}"],
        ["Original LFW images", f"{CONFIG['original_images']:,}"],
        ["Evaluation pairs", f"{CONFIG['subset_pairs']} total ({CONFIG['genuine_pairs']} genuine, {CONFIG['impostor_pairs']} impostor)"],
        ["Unique images used for filtering", f"{CONFIG['unique_pair_images']}"],
        ["Generated filtered images", f"{CONFIG['filtered_images']:,}"],
        ["Filter categories", "blur, brightness, skin_smooth, eye_enlarge, face_slim, color_tone"],
    ], [5.0 * cm, 11.0 * cm])
    p(
        "Ethically, LFW is a public research benchmark, but face images remain sensitive biometric data. "
        "This system is an academic prototype and should not be used for surveillance, access control, or production identity decisions. "
        "The generated filtered images are derived from LFW and should follow the same research-use expectations."
    )

    h1("5. Implementation Details")
    p(
        "The implementation uses Python, OpenCV, PIL, NumPy, pandas, scikit-learn, matplotlib, seaborn, DeepFace, and Gradio. "
        "The filter functions are deterministic and preserve the original LFW identity-folder structure. "
        "Embedding extraction uses caching so repeated runs do not recompute embeddings unnecessarily. "
        "The Gradio app lets a user upload two face images, select a recognition model and mitigation mode, apply a filter, and receive a similarity score and verification decision."
    )
    h2("Filter Simulation")
    bullet("blur: Gaussian blur blended with the original image.")
    bullet("brightness: brightness and saturation enhancement.")
    bullet("skin_smooth: HSV skin mask with bilateral filtering.")
    bullet("eye_enlarge: subtle eye-region warp with landmark or OpenCV fallback.")
    bullet("face_slim: local horizontal compression of the detected face region.")
    bullet("color_tone: warm red-blue channel shift with vignette.")
    h2("Novel Improvement Added")
    p(
        "The project adds two improvements beyond the paper-style baseline. First, it introduces an improved-backbone path intended for AdaFace, with GhostFaceNet fallback under the current environment. "
        "Second, it replaces simple mean-shift correction with a per-filter Ridge-regression linear corrector that learns a mapping from filtered embeddings back toward original embeddings. "
        "This correction is trained with an identity-based split to reduce leakage between training and evaluation identities."
    )

    h1("6. Experiments and Results")
    h2("6.1 ArcFace Baseline")
    p(
        f"The ArcFace baseline achieved {pct(summary['arc_original_acc'])} accuracy on original images and an average filtered accuracy of {pct(summary['arc_filter_acc'])}. "
        f"The average filtered FNMR@FMR=1% was {pct(summary['arc_filter_fnmr'])}."
    )
    table(recognition_rows(results["arcface"]), [3.3 * cm, 3.3 * cm, 4.5 * cm, 4.5 * cm])
    plot("accuracy_per_filter.png", "Figure 1: ArcFace accuracy per filter, generated from the submitted evaluation metrics.")

    h2("6.2 Improved Backbone Path")
    p(
        f"The improved-backbone path achieved {pct(summary['imp_original_acc'])} original-image accuracy and {pct(summary['imp_filter_acc'])} average filtered accuracy. "
        f"The average filtered FNMR@FMR=1% was {pct(summary['imp_filter_fnmr'])}, a {summary['improved_reduction']:.1f}% average reduction relative to the ArcFace filter baseline."
    )
    table(recognition_rows(results["improved"]), [3.3 * cm, 3.3 * cm, 4.5 * cm, 4.5 * cm])
    plot("comparison/accuracy_4way_comparison.png", "Figure 2: Four-way comparison of baseline, mean-shift, improved backbone, and linear correction.")

    h2("6.3 Mitigation Comparison")
    p(
        f"The mitigation comparison shows that mean-shift reduces average FNMR relative to the ArcFace baseline, while the improved-backbone plus linear-corrector condition has an average FNMR@FMR=1% of {pct(summary['linear_fnmr'])}. "
        f"Across filters, the combined condition reduces average FNMR by {summary['combined_reduction']:.1f}% relative to the paper-style ArcFace baseline. "
        "However, the linear corrector is not uniformly better for every filter; this is discussed as a limitation because the prototype uses a small calibration subset."
    )
    table(comparison_rows(results["summary"]), [3.1 * cm, 3.2 * cm, 3.2 * cm, 3.7 * cm, 3.7 * cm])
    plot("comparison/fnmr_improvement.png", "Figure 3: FNMR@FMR=1% comparison across filters and mitigation settings.")
    plot("comparison/det_curves_comparison.png", "Figure 4: DET curves for the selected disruptive filter comparison.")
    plot("comparison/summary_table.png", "Figure 5: Summary table visualization generated from the comparison CSV.")

    h1("7. Discussion and Limitations")
    p(
        "The results confirm the usefulness of an end-to-end evaluation framework: different filters produce different shifts in similarity distributions, and model choice has a measurable effect on recognition behavior. "
        "The improved-backbone path performs better than the ArcFace baseline on this submitted LFW subset. "
        "Mitigation is more complicated: while average FNMR improves in the combined condition, the per-filter linear corrector is mixed and can underperform for some filters, indicating sensitivity to calibration size and data split."
    )
    bullet("The original paper used FRGC and real mobile AR filters; this project uses public LFW data and programmatic filter simulation.")
    bullet("The evaluation subset is intentionally limited to 300 genuine and 300 impostor pairs for runtime practicality.")
    bullet("AdaFace could not be installed in the Python 3.13/macOS environment, so GhostFaceNet is used as the configured fallback.")
    bullet("The linear corrector needs original-filtered calibration pairs and may not generalize to unseen filters.")
    bullet("The prototype does not evaluate demographic fairness, real app filters, or large-scale deployment behavior.")

    h1("8. Conclusion and Future Work")
    p(
        "This project successfully implements a full computer vision pipeline for studying the effect of facial filters on face recognition. "
        "It prepares a dataset, generates filtered images, extracts embeddings, evaluates recognition degradation, applies mitigation, produces plots, and deploys a live Gradio demo. "
        "The implementation follows the core methodology of the selected research paper while adding a practical model-comparison and linear-correction extension. "
        "Future work should evaluate the original FRGC dataset, apply real mobile AR filters through an emulator, train a filter-agnostic corrector, test larger pair subsets, and study fairness across demographic attributes."
    )

    h1("9. References")
    refs = [
        "[1] K. Ozturk, L. Conwill, J. Gutierrez, K. Bowyer, and W. J. Scheirer, A Comprehensive Evaluation Framework for the Study of the Effects of Facial Filters on Face Recognition Accuracy, IJCB 2025 / arXiv:2507.17729.",
        "[2] J. Deng, J. Guo, N. Xue, and S. Zafeiriou, ArcFace: Additive Angular Margin Loss for Deep Face Recognition, CVPR 2019.",
        "[3] M. Kim, A. K. Jain, and X. Liu, AdaFace: Quality Adaptive Margin for Face Recognition, CVPR 2022.",
        "[4] G. B. Huang, M. Ramesh, T. Berg, and E. Learned-Miller, Labeled Faces in the Wild: A Database for Studying Face Recognition in Unconstrained Environments, UMass Technical Report 07-49, 2007.",
        "[5] P. Riccio et al., OpenFilter: A Framework to Democratize Research Access to Social Media AR Filters, NeurIPS 2022.",
    ]
    for ref in refs:
        story.append(Paragraph(ref, styles["BodyJustify"]))

    doc.build(story)


def set_docx_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def build_docx(results: dict) -> None:
    summary = metrics_summary(results)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [("Heading 1", 16, "185FA5"), ("Heading 2", 13, "0C447C")]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "Facial Filter Impact on Face Recognition - Final Project Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "Computer Vision | FAST School of Computing | Spring 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def p(text: str) -> None:
        para = doc.add_paragraph(text)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def bullet(text: str) -> None:
        para = doc.add_paragraph(text, style="List Bullet")
        para.paragraph_format.space_after = Pt(4)

    def add_table(rows: list[list[str]]) -> None:
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                cell = t.cell(r, c)
                cell.text = value
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(2)
                    for run in para.runs:
                        run.font.size = Pt(9)
                if r == 0:
                    set_docx_cell_shading(cell, "E6F1FB")
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(CONFIG["project_title"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph("Final Project Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string("1F4D78")
    doc.add_paragraph(CONFIG["institution"]).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(CONFIG["course"]).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Based on: {CONFIG['paper']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    sections = [
        ("Abstract", [
            f"This project implements an end-to-end computer vision system for evaluating how facial filters affect face verification accuracy. It reimplements the core methodology of Ozturk et al. using LFW, six simulated facial-filter categories, embedding extraction, biometric evaluation, mitigation, plots, and a live Gradio demo. On the submitted 600-pair subset, ArcFace achieved {pct(summary['arc_filter_acc'])} average filtered accuracy, while the improved-backbone path achieved {pct(summary['imp_filter_acc'])}. The improved backbone reduced average FNMR@FMR=1% by {summary['improved_reduction']:.1f}% relative to the ArcFace baseline."
        ]),
        ("1. Introduction and Problem Definition", [
            "Face recognition systems are increasingly exposed to images modified by camera and social-media filters. These filters can change color, texture, illumination, and local facial geometry, shifting the embedding produced by a recognition model. The project problem is to measure and mitigate the effect of such filters on face verification accuracy."
        ]),
        ("2. Literature Review", [
            "Ozturk et al. provide the primary evaluation framework: controlled data, principled filter selection, FNMR/FMR metrics, and DET curves. ArcFace provides a strong angular-margin face-recognition baseline. AdaFace motivates quality-aware recognition for degraded face images. LFW provides the public verification benchmark used in this course implementation."
        ]),
        ("3. Methodology and System Architecture", [
            "The system pipeline is dataset setup -> filter simulation -> embedding extraction -> verification scoring -> mitigation -> visualization -> Gradio deployment. The implementation is modular, with each stage available as an independent script."
        ]),
        ("4. Dataset Description", [
            f"The prepared LFW dataset contains {CONFIG['original_images']:,} original images across {CONFIG['identities']:,} identities. The submitted subset contains {CONFIG['genuine_pairs']} genuine and {CONFIG['impostor_pairs']} impostor pairs. The filter scripts generate {CONFIG['filtered_images']:,} filtered images across six categories."
        ]),
        ("5. Implementation Details", [
            "The implementation uses Python, OpenCV, PIL, NumPy, pandas, scikit-learn, matplotlib, seaborn, DeepFace, and Gradio. It includes ArcFace baseline extraction, an improved-backbone path with GhostFaceNet fallback when AdaFace is unavailable, mean-shift mitigation, and per-filter Ridge linear correction."
        ]),
        ("6. Experiments and Results", [
            f"ArcFace original accuracy was {pct(summary['arc_original_acc'])}; improved-backbone original accuracy was {pct(summary['imp_original_acc'])}. Average filtered FNMR@FMR=1% decreased from {pct(summary['arc_filter_fnmr'])} with ArcFace to {pct(summary['imp_filter_fnmr'])} with the improved backbone. The combined improved-backbone plus linear-corrector condition reduced average FNMR by {summary['combined_reduction']:.1f}% relative to the ArcFace baseline, although per-filter behavior was mixed."
        ]),
        ("7. Discussion and Limitations", [
            "The results support the need for filter-specific evaluation. The improved backbone is stronger on this subset, but the linear corrector is not uniformly better across filters. Limitations include the use of LFW rather than FRGC, programmatic filters rather than real app filters, a 600-pair subset for runtime, and GhostFaceNet fallback because AdaFace was unavailable in the environment."
        ]),
        ("8. Conclusion and Future Work", [
            "The project successfully implements an end-to-end facial-filter recognition evaluation system. Future work should test FRGC, use real AR filters through an Android emulator, evaluate larger subsets, train filter-agnostic correction, and study demographic fairness."
        ]),
    ]
    for heading, paras in sections:
        doc.add_heading(heading, level=1)
        for text in paras:
            p(text)
        if heading == "3. Methodology and System Architecture":
            add_table([
                ["Stage", "Script / Module", "Purpose"],
                ["Dataset setup", "src/01_setup_dataset.py", "Prepare LFW and verification pairs."],
                ["Filtering", "src/02_apply_filters.py", "Generate six simulated filter categories."],
                ["Embedding", "src/04_extract_embeddings.py", "Extract and cache face embeddings."],
                ["Evaluation", "src/05_evaluate_recognition.py", "Compute accuracy, FNMR, FMR, and DET curves."],
                ["Demo", "app/app.py", "Run the Gradio verification dashboard."],
            ])
        if heading == "4. Dataset Description":
            add_table([
                ["Dataset Item", "Value"],
                ["Original LFW identities", f"{CONFIG['identities']:,}"],
                ["Original LFW images", f"{CONFIG['original_images']:,}"],
                ["Evaluation pairs", f"{CONFIG['subset_pairs']} total"],
                ["Generated filtered images", f"{CONFIG['filtered_images']:,}"],
                ["Filter categories", "blur, brightness, skin_smooth, eye_enlarge, face_slim, color_tone"],
            ])
        if heading == "6. Experiments and Results":
            doc.add_heading("ArcFace baseline metrics", level=2)
            add_table(recognition_rows(results["arcface"]))
            doc.add_heading("Improved-backbone metrics", level=2)
            add_table(recognition_rows(results["improved"]))
            doc.add_heading("Mitigation comparison", level=2)
            add_table(comparison_rows(results["summary"]))
            for rel, caption in [
                ("accuracy_per_filter.png", "ArcFace accuracy per filter."),
                ("comparison/accuracy_4way_comparison.png", "Four-way comparison plot."),
                ("comparison/fnmr_improvement.png", "FNMR comparison plot."),
            ]:
                path = PLOTS_DIR / rel
                if path.exists():
                    doc.add_picture(str(path), width=Inches(6.2))
                    cap = doc.add_paragraph(caption)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("9. References", level=1)
    for ref in [
        "[1] K. Ozturk, L. Conwill, J. Gutierrez, K. Bowyer, and W. J. Scheirer, A Comprehensive Evaluation Framework for the Study of the Effects of Facial Filters on Face Recognition Accuracy, IJCB 2025 / arXiv:2507.17729.",
        "[2] J. Deng, J. Guo, N. Xue, and S. Zafeiriou, ArcFace: Additive Angular Margin Loss for Deep Face Recognition, CVPR 2019.",
        "[3] M. Kim, A. K. Jain, and X. Liu, AdaFace: Quality Adaptive Margin for Face Recognition, CVPR 2022.",
        "[4] G. B. Huang, M. Ramesh, T. Berg, and E. Learned-Miller, Labeled Faces in the Wild: A Database for Studying Face Recognition in Unconstrained Environments, UMass Technical Report 07-49, 2007.",
        "[5] P. Riccio et al., OpenFilter: A Framework to Democratize Research Access to Social Media AR Filters, NeurIPS 2022.",
    ]:
        p(ref)

    doc.core_properties.title = CONFIG["project_title"] + " - Final Project Report"
    doc.core_properties.author = "Computer Vision Course Project Group"
    doc.save(OUTPUT_DOCX)


def main() -> None:
    results = load_results()
    build_pdf(results)
    build_docx(results)
    print(f"PDF report saved to {OUTPUT_PDF}")
    print(f"DOCX report saved to {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
